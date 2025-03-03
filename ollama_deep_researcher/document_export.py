from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
import re
import os
from slugify import slugify
from datetime import datetime

def extract_sections(markdown_content):
    """Extract sections from markdown content"""
    # Pattern to match headers and content
    sections = {}
    current_section = None
    lines = markdown_content.split('\n')
    
    for line in lines:
        if line.startswith('# '):
            current_section = line[2:].strip()
            sections[current_section] = []
        elif line.startswith('## '):
            subsection = line[3:].strip()
            if current_section:
                sections[current_section].append(f"### {subsection}")
        elif current_section:
            sections[current_section].append(line)
    
    return {k: '\n'.join(v) for k, v in sections.items()}

def extract_citations(markdown_content):
    """Extract citations from markdown content"""
    citation_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    citations = re.findall(citation_pattern, markdown_content)
    return citations

def export_to_word(markdown_content, topic="Research Topic", output_path=None):
    """Convert markdown research output to formatted Word document"""
    if output_path is None:
        slug = slugify(topic[:30])
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        output_path = f"{slug}-{timestamp}.docx"
    
    doc = Document()
    
    # Document style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Add title
    title = doc.add_heading(f"Research Summary: {topic}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(datetime.now().strftime("%B %d, %Y")).italic = True
    
    doc.add_paragraph()  # Add spacing
    
    # Process sections
    sections = extract_sections(markdown_content)
    
    # Main content
    for section_title, section_content in sections.items():
        # Add section header
        doc.add_heading(section_title, 1)
        
        # Process markdown content
        p = doc.add_paragraph()
        # Simple markdown conversion for paragraphs
        for line in section_content.split('\n'):
            if line.startswith('### '):
                doc.add_heading(line[4:], 2)
            elif line.strip():
                p = doc.add_paragraph()
                p.add_run(line)
    
    # Add sources section
    doc.add_heading("Sources", 1)
    citations = extract_citations(markdown_content)
    
    for i, (text, url) in enumerate(citations):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {text}: ").bold = True
        p.add_run(url)
    
    # Save document
    doc.save(output_path)
    print(f"Research exported to: {os.path.abspath(output_path)}")
    return os.path.abspath(output_path)